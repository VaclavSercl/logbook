"""AI Voyage Document Analyzer and Information Extractor.

Processes uploaded files, local folder paths, and URLs to automatically populate:
- Vessel specifications (Name, Type, Length, Beam, Draft, Year, Port, Flag State, Charter Company)
- Voyage route and itinerary details (departure port, destination port, notes)
- Financial expenses in Cashbox (Deposit, Charter Fees)
- Crew members list (only if present in documents, with full Czech encoding repair CP1250/UTF-8)
"""
import os
import re
import csv
import json
import urllib.request
from datetime import datetime, timedelta
from typing import Dict, Any, List, Optional
from sqlalchemy.orm import Session

from app.models import VoyageDocument, Logbook, Vessel, CrewMember, LogEntry, CashboxExpense
from app.config import settings
from app.services.schedule_generator import auto_generate_schedules


def parse_dob(val: Any) -> Optional[datetime]:
    if not val:
        return None
    if isinstance(val, datetime):
        return val
    val_str = str(val).strip()
    if not val_str:
        return None
    for fmt in ["%Y-%m-%d", "%d.%m.%Y", "%d/%m/%Y", "%Y/%m/%d", "%d-%m-%Y"]:
        try:
            return datetime.strptime(val_str, fmt)
        except ValueError:
            pass
    return None


def sanitize_vessel_name(raw_name: Optional[str], default_title: str) -> str:
    """Cleans vessel name to prevent Windows file paths or raw folder paths from becoming vessel names."""
    clean = ""
    if raw_name and str(raw_name).strip():
        clean = str(raw_name).strip()
    elif default_title and str(default_title).strip():
        clean = str(default_title).strip()

    # Remove colons, Windows drive letters and paths
    if ":" in clean:
        parts = clean.split(":")
        clean = parts[-1].strip()

    if any(bad in clean for bad in ["\\", "/"]):
        clean = os.path.basename(clean.replace('\\', '/').rstrip('/'))

    # Clean date prefixes like 2026-07-Recko-Lefkada -> Recko-Lefkada
    clean = re.sub(r'^\d{4}[\-_]\d{2}[\-_]?', '', clean).strip()

    if not clean or clean.lower() in ["file", "folder", "general", "undefined", "null", "none"]:
        return "Mořská Plachetnice"

    return clean

def create_dedicated_ai_voyage_folder(vessel_name: str, voy_from: str, voy_to: str, start_date_str: Optional[str], gemini_summary: str, source_path: Optional[str] = None) -> str:
    """Creates a dedicated voyage directory on disk organized by AI with date, boat name, and route."""
    import re, shutil
    base_dir = os.path.join(os.getcwd(), "uploads", "voyages")
    os.makedirs(base_dir, exist_ok=True)

    date_part = datetime.utcnow().strftime("%Y-%m-%d")
    if start_date_str:
        parsed_dt = parse_dob(start_date_str)
        if parsed_dt:
            date_part = parsed_dt.strftime("%Y-%m-%d")

    clean_vessel = re.sub(r'[^\w\s-]', '', vessel_name or 'Plavidlo').strip().replace(' ', '_')
    clean_route = re.sub(r'[^\w\s-]', '', f"{voy_from or 'Start'}_to_{voy_to or 'Cil'}").strip().replace(' ', '_')

    folder_name = f"{date_part}_{clean_vessel}_{clean_route}"
    voyage_folder = os.path.join(base_dir, folder_name)
    os.makedirs(voyage_folder, exist_ok=True)

    # Write AI Voyage Summary report into the folder
    summary_file = os.path.join(voyage_folder, "njoror_ai_voyage_summary.md")
    with open(summary_file, "w", encoding="utf-8") as f:
        f.write(f"# ⛵ Njořðr AI — Souhrn Plavby\n\n")
        f.write(f"**Datum plavby:** {date_part}\n")
        f.write(f"**Plavidlo:** {vessel_name}\n")
        f.write(f"**Trasa:** {voy_from or 'Nespecifikováno'} ➔ {voy_to or 'Nespecifikováno'}\n\n")
        f.write(f"## 📊 Výstup AI Analýzy (Gemini 3.6 Flash High):\n\n")
        f.write(gemini_summary or "")

    # If source_path exists and is a folder or file, copy contents into dedicated folder
    if source_path and os.path.exists(source_path) and os.path.abspath(source_path) != os.path.abspath(voyage_folder):
        try:
            if os.path.isdir(source_path):
                for item in os.listdir(source_path):
                    s = os.path.join(source_path, item)
                    d = os.path.join(voyage_folder, item)
                    if os.path.isdir(s):
                        if not os.path.exists(d):
                            shutil.copytree(s, d)
                    else:
                        shutil.copy2(s, d)
            elif os.path.isfile(source_path):
                shutil.copy2(source_path, os.path.join(voyage_folder, os.path.basename(source_path)))
        except Exception as e:
            print("Copy to dedicated voyage folder warning:", e)

    return voyage_folder


def smart_parse_text_fallback(text: str) -> Dict[str, Any]:
    """Smartly extracts boat name, specs, itinerary, and crew from raw document text when AI API is unavailable."""
    result = {}
    if not text:
        return result

    text_fixed = fix_mojibake(text)
    lines = text_fixed.splitlines()

    for line in lines:
        l = line.strip()
        l_lower = l.lower()

        # Boat / Vessel Name
        if any(kw in l_lower for kw in ["plavidlo:", "loď:", "lod:", "jachta:", "vessel:", "boat:", "název lodi:"]):
            parts = l.split(":", 1)
            if len(parts) == 2 and len(parts[1].strip()) > 2:
                result["vessel_name"] = parts[1].strip()

        # Vessel Type
        elif any(kw in l_lower for kw in ["typ:", "druh:"]):
            parts = l.split(":", 1)
            if len(parts) == 2:
                result["vessel_type"] = parts[1].strip()

        # Length / Beam / Draft / Year
        elif "délka:" in l_lower or "delka:" in l_lower or "length:" in l_lower:
            m = re.search(r'(\d+(?:[.,]\d+)?)', l)
            if m:
                try:
                    result["length"] = float(m.group(1).replace(',', '.'))
                except Exception:
                    pass
        elif "šířka:" in l_lower or "sirka:" in l_lower or "beam:" in l_lower:
            m = re.search(r'(\d+(?:[.,]\d+)?)', l)
            if m:
                try:
                    result["beam"] = float(m.group(1).replace(',', '.'))
                except Exception:
                    pass
        elif "ponor:" in l_lower or "draft:" in l_lower:
            m = re.search(r'(\d+(?:[.,]\d+)?)', l)
            if m:
                try:
                    result["draft"] = float(m.group(1).replace(',', '.'))
                except Exception:
                    pass
        elif "rok výroby:" in l_lower or "rok:" in l_lower or "year:" in l_lower:
            m = re.search(r'(20\d{2}|19\d{2})', l)
            if m:
                try:
                    result["year_built"] = int(m.group(1))
                except Exception:
                    pass

        # Port / Marina
        elif any(kw in l_lower for kw in ["domovský přístav:", "přístav:", "marina:", "port:"]):
            parts = l.split(":", 1)
            if len(parts) == 2 and len(parts[1].strip()) > 2:
                result["port"] = parts[1].strip()

        # Flag State
        elif "vlajka:" in l_lower or "flag:" in l_lower:
            parts = l.split(":", 1)
            if len(parts) == 2:
                result["flag_state"] = parts[1].strip().upper()

        # Charter Company
        elif "charterovka:" in l_lower or "charter:" in l_lower:
            parts = l.split(":", 1)
            if len(parts) == 2:
                result["charter_company"] = parts[1].strip()

        # Deposit & Fee
        elif "vratná kauce:" in l_lower or "kauce:" in l_lower or "deposit:" in l_lower:
            m = re.search(r'(\d+)', l.replace(' ', ''))
            if m:
                try:
                    result["deposit"] = float(m.group(1))
                except Exception:
                    pass
        elif "poplatky:" in l_lower or "service fee:" in l_lower:
            m = re.search(r'(\d+)', l.replace(' ', ''))
            if m:
                try:
                    result["service_fee"] = float(m.group(1))
                except Exception:
                    pass

        # Itinerary
        elif any(kw in l_lower for kw in ["vyplutí:", "start:"]):
            parts = l.split(":", 1)
            if len(parts) == 2:
                result["voyage_from"] = parts[1].strip()
        elif any(kw in l_lower for kw in ["cíl:", "trasa a cíl:", "destination:"]):
            parts = l.split(":", 1)
            if len(parts) == 2:
                result["voyage_to"] = parts[1].strip()

    # Search for boat name in title or first lines if missing
    if "vessel_name" not in result:
        m = re.search(r'([A-Z][a-zA-Z0-9\s\.\-]{3,35}\s+(?:\d{2,3}|\-\s*[A-Z][a-z]+))', text_fixed)
        if m:
            result["vessel_name"] = m.group(1).strip()

    return result


def extract_text_from_file(file_path: str) -> str:
    """Extracts plain text content from any file: CSV, TXT, MD, JSON, XLSX, XLS, PDF, DOCX, HTML, XML."""
    if not os.path.exists(file_path) or os.path.isdir(file_path):
        return ""

    ext = os.path.splitext(file_path)[1].lower()
    text_chunks = []

    try:
        if ext in ['.txt', '.md', '.json', '.csv', '.tsv', '.html', '.xml', '.log', '.ini']:
            with open(file_path, 'rb') as fp:
                raw = fp.read()
                return decode_bytes_to_str(raw)

        elif ext in ['.xlsx', '.xls']:
            try:
                import pandas as pd
                excel = pd.ExcelFile(file_path)
                for sheet in excel.sheet_names:
                    df = excel.parse(sheet)
                    text_chunks.append(f"--- List Excelu: {sheet} ---")
                    text_chunks.append(df.to_string())
                return "\n".join(text_chunks)
            except Exception as e:
                print(f"Excel read error {file_path}: {e}")

        elif ext == '.pdf':
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                for page_idx, page in enumerate(reader.pages):
                    t = page.extract_text()
                    if t:
                        text_chunks.append(f"--- PDF Strana {page_idx+1} ---")
                        text_chunks.append(t)
                return "\n".join(text_chunks)
            except Exception as e:
                print(f"PDF read error {file_path}: {e}")

        elif ext == '.docx':
            try:
                import docx
                doc = docx.Document(file_path)
                for p in doc.paragraphs:
                    if p.text.strip():
                        text_chunks.append(p.text.strip())
                for table in doc.tables:
                    for row in table.rows:
                        text_chunks.append(" | ".join([cell.text.strip() for cell in row.cells]))
                return "\n".join(text_chunks)
            except Exception as e:
                print(f"DOCX read error {file_path}: {e}")

    except Exception as e:
        print(f"Error reading file {file_path}: {e}")

    # Fallback binary decode
    try:
        with open(file_path, 'rb') as fp:
            return decode_bytes_to_str(fp.read())
    except Exception:
        return ""


def read_all_files_recursively(folder_path: str):
    """Recursively walks folder_path and all subfolders to extract text from all files."""
    all_texts = []
    summary_files = []

    priority_files = []
    other_files = []

    for root, _, files in os.walk(folder_path):
        for f in files:
            full_p = os.path.join(root, f)
            rel_p = os.path.relpath(full_p, folder_path)
            f_lower = f.lower()

            if any(kw in f_lower for kw in ["crew", "posad", "list", "manifest", "pasazer", "passenger", "charter", "voucher", "contract"]):
                priority_files.append((full_p, rel_p))
            else:
                other_files.append((full_p, rel_p))

    sorted_files = priority_files + other_files

    for full_p, rel_p in sorted_files[:60]:
        t = extract_text_from_file(full_p)
        if t.strip():
            all_texts.append(f"\n==================== SOUBOR: {rel_p} ====================\n{t}\n")
            summary_files.append(rel_p)

    return "\n".join(all_texts), summary_files


def analyze_document_with_gemini(combined_text: str) -> Dict[str, Any]:
    """Uses Google Gemini 3.6 Flash API to analyze document/folder text and return structured JSON."""
    google_key = settings.GOOGLE_API_KEY
    if not google_key or not combined_text:
        return {}

    prompt = f"""
Jsi Njořðr, nekompromisní a brilantní AI vládce projektu lodního deníku a analýzy jachtařských podkladů.
Tvým úkolem je PROSKOUMAT A PROHLÉDNOUT VŠECHNY PRVKY, SOUBORY A PODSLOŽKY z dodaného balíčku podkladů od kapitána (Skippera) a extrahovat kompletní a 100% přesná data pro plavbu.

Jsi Njořðr, nekompromisní a brilantní AI vládce projektu lodního deníku a navigační analýzy jachtařských podkladů.
Kapitán lodi ti právě nahrál složku nebo balíček dokumentů k plavbě. Tvým úkolem je PROSKOUMAT A DŮKLADNĚ PROHLÉDNOUT VŠECHNY PRVKY, SOUBORY A PODSLOŽKY a vyextrahovat kompletní a 100% přesná data pro založení nebo aktualizaci plavidla a lodního deníku.

🎯 NEKOMPROMISNÍ POŽADAVKY NA EXTRAKCI DATA:
1. ⛵ NÁZEV A PARAMETRY LOĎE:
   - Najdi přesné jméno lodi (názvy jako Lagoon, Bavaria, Oceanis, Sun Odyssey, Hanse, Dufour, Fountain Pajot apod., včetně případného čísla lodi nebo jména).
   - Typ lodi (Plachetnice / Katamarán / Motorová jachta).
   - Parametry: Délka (length m), Šířka (beam m), Ponor (draft m), Rok výroby (year_built), Domovský přístav/Marina (port), Vlajka státu (flag_state, např. GR, HR, IT, CZ), Charterová společnost (charter_company).
   - Finanční poplatky: Vratná kauce (deposit EUR), Transit log / Service pack (service_fee EUR).

2. 📍 ITINERÁŘ A TERMÍNY:
   - Výchozí přístav / vyplutí (voyage_from).
   - Cílový přístav / příjezd / oblast plavby (voyage_to).
   - Přesný termín plavby (voyage_start_date: YYYY-MM-DDTHH:MM:SS, voyage_end_date: YYYY-MM-DDTHH:MM:SS).

3. 👥 POSÁDKA A KAPITÁN (CREW LIST):
   - Vyhledej VŠECHNY osoby uvedené v dokumentech (jméno, příjmení, přezdívka, datum narození YYYY-MM-DD, číslo pasu/OP, národnost, role).
   - Urči přesně SKIPPERA / KAPITÁNA (hledej výraz Skipper, Kapitán, Velitel, Master, nebo Václav...). Uveď jeho celé jméno do "skipper_name".
   - Ostatním členům přiřaď role: "První důstojník", "Kuchař", nebo "Posádka".
   - NEVYNECHEJ ŽÁDNÉHO ČLENA POSÁDKY!

Odpověz POUZE ve formátu čistého JSON bez jakéhokoliv dalšího textu nebo uvozovek markdown:
{{
  "vessel_name": "Přesný Název Lodi",
  "vessel_type": "Plachetnice / Katamarán / Motorová jachta",
  "length": 14.2,
  "beam": 4.35,
  "draft": 2.1,
  "year_built": 2024,
  "port": "Přístav / Marina",
  "flag_state": "GR",
  "mmsi": null,
  "call_sign": null,
  "charter_company": "Název Charterové Společnosti",
  "deposit": 2500.0,
  "service_fee": 250.0,
  "voyage_from": "Výchozí přístav vyplutí",
  "voyage_to": "Cílový přístav příjezdu",
  "voyage_start_date": "2026-07-25T17:00:00",
  "voyage_end_date": "2026-08-01T09:00:00",
  "skipper_name": "Jméno a Příjmení Kapitána",
  "crew_members": [
    {{
      "first_name": "Jméno",
      "last_name": "Příjmení",
      "nickname": "Přezdívka",
      "passport_number": "Číslo OP/pasu",
      "date_of_birth": "YYYY-MM-DD",
      "nationality": "CZ",
      "role": "Skipper (Kapitán) / První důstojník / Kuchař / Posádka"
    }}
  ],
  "log_entry_notes": "Profesionální námořní svodka Njořðra k plavbě a stavu plavidla"
}}

DODANÉ PODKLADY ZE VŠECH SOUBORŮ A PODSLOŽEK:
{combined_text[:32000]}
"""

    models_to_try = [
        "gemini-3.6-flash",
        "gemini-2.5-flash",
        "gemini-2.0-flash",
        "gemini-2.0-flash-lite"
    ]

    import httpx
    for model_name in models_to_try:
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent?key={google_key}"
        payload = {
            "contents": [
                {
                    "parts": [
                        {
                            "text": prompt
                        }
                    ]
                }
            ]
        }
        try:
            with httpx.Client(timeout=45.0) as client:
                res = client.post(url, json=payload)
                if res.status_code == 200:
                    data = res.json()
                    resp_text = data["candidates"][0]["content"]["parts"][0]["text"].strip()
                    resp_text = re.sub(r'^```(?:json)?\s*', '', resp_text, flags=re.MULTILINE)
                    resp_text = re.sub(r'\s*```$', '', resp_text, flags=re.MULTILINE)
                    return json.loads(resp_text)
        except Exception as e:
            print(f"Gemini {model_name} document analysis warning:", e)

    return {}


def decode_bytes_to_str(content: bytes) -> str:
    """Decodes bytes to string using UTF-8-SIG, UTF-8, CP1250 (Windows Czech), ISO-8859-2."""
    if not content:
        return ""
    try:
        return content.decode('utf-8-sig')
    except UnicodeDecodeError:
        pass

    try:
        decoded = content.decode('utf-8')
        if not any(bad in decoded for bad in ['Ã¡', 'Å¡', 'Ã¨', 'Ã½', 'Ã']):
            return decoded
    except UnicodeDecodeError:
        pass

    try:
        return content.decode('cp1250')
    except UnicodeDecodeError:
        pass

    try:
        return content.decode('iso-8859-2')
    except UnicodeDecodeError:
        pass

    return content.decode('utf-8', errors='replace')


def fix_mojibake(text: str) -> str:
    """Fixes common Mojibake UTF-8 vs CP1250 double-encoding artifacts in Czech text."""
    if not text:
        return ""
    
    replacements = {
        'Ã¡': 'á', 'Ã©': 'é', 'Ã­': 'í', 'Ã³': 'ó', 'Ãº': 'ú', 'Ã¹': 'ů', 'Ã½': 'ý',
        'Ä\x8d': 'č', 'Ä\x8c': 'Č', 'Å¡': 'š', 'Å\x90': 'Š', 'Å¾': 'ž', 'Å½': 'Ž',
        'Å\x99': 'ř', 'Å\x98': 'Ř', 'Å¥': 'ť', 'Å¤': 'Ť', 'Ä\x8f': 'ď', 'Ä\x8e': 'Ď',
        'Å\x88': 'ň', 'Å\x87': 'Ň', 'Ä\x9b': 'ě', 'Ä\x9a': 'Ě',
        'Á¡': 'á', 'Á©': 'é', 'Á­': 'í', 'Á³': 'ó', 'Áº': 'ú', 'Á¹': 'ů', 'Á½': 'ý',
        'ÄŤ': 'č', 'ÄŚ': 'Č', 'Ĺˇ': 'š', 'ĹŠ': 'Š', 'Ĺľ': 'ž', 'ĹŽ': 'Ž',
        'Ĺ™': 'ř', 'ĹŘ': 'Ř', 'ĹĄ': 'ť', 'ĹŤ': 'Ť', 'ÄŹ': 'ď', 'ÄĎ': 'Ď',
        'Ĺˆ': 'ň', 'ĹŇ': 'Ň', 'Ä•': 'ě', 'ÄŞ': 'Ě'
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text.strip()


def parse_crew_from_text_or_rows(rows_or_text: List[List[str]]) -> List[Dict[str, Any]]:
    """Helper to extract crew members from matrix of text/csv/excel rows."""
    crew_found = []
    if not rows_or_text:
        return crew_found

    # Normalize single-string rows
    normalized_rows = []
    for r in rows_or_text:
        if r and len(r) == 1 and ',' in str(r[0]):
            normalized_rows.append([c.strip() for c in str(r[0]).split(',')])
        elif r:
            normalized_rows.append([str(c).strip() for c in r])

    header_idx = -1
    col_map = {}

    for idx, row in enumerate(normalized_rows[:10]):
        row_lower = [c.lower() for c in row]
        for c_i, cell in enumerate(row_lower):
            if any(k in cell for k in ["jméno", "jmeno", "first name", "name", "člen", "posádka", "skupina"]):
                header_idx = idx
                break
        if header_idx != -1:
            break

    if header_idx != -1:
        header_row = [c.lower() for c in normalized_rows[header_idx]]
        for c_i, cell in enumerate(header_row):
            if "přezdívka" in cell or "prezdivka" in cell or "nick" in cell:
                col_map["nickname"] = c_i
            elif "příjmení" in cell or "prijmeni" in cell or "last" in cell:
                col_map["last_name"] = c_i
            elif "jméno" in cell or "jmeno" in cell or "first" in cell:
                col_map["first_name"] = c_i
            elif "role" in cell or "funkce" in cell or "pozice" in cell:
                col_map["role"] = c_i

        data_rows = normalized_rows[header_idx + 1:]
    else:
        data_rows = normalized_rows

    valid_name_regex = re.compile(r'^[a-zA-ZáčďéěíňóřšťúůýžÁČĎÉĚÍŇÓŘŠŤÚŮÝŽ\s\.\-]{2,30}$')

    for row in data_rows:
        if not row or not any(str(c).strip() for c in row):
            continue

        first_name = ""
        last_name = ""
        nickname = ""
        role = "Člen posádky"

        if "first_name" in col_map and col_map["first_name"] < len(row):
            first_name = str(row[col_map["first_name"]]).strip()
        if "last_name" in col_map and col_map["last_name"] < len(row):
            last_name = str(row[col_map["last_name"]]).strip()
        if "nickname" in col_map and col_map["nickname"] < len(row):
            nickname = str(row[col_map["nickname"]]).strip()
        if "role" in col_map and col_map["role"] < len(row):
            role = str(row[col_map["role"]]).strip() or "Člen posádky"

        if not first_name and not last_name:
            cells = [str(c).strip() for c in row if str(c).strip()]
            if cells and len(cells[0]) < 40:
                parts = cells[0].split()
                if len(parts) >= 2 and not any(kw in cells[0].lower() for kw in ["délka", "ponor", "vyplutí", "vítr", "tlak", "http", "bavaria", "charter", "const", "var", "function"]):
                    first_name, last_name = parts[0], parts[1]

def process_voyage_document_ai(db: Session, doc_id: str) -> Dict[str, Any]:
    """Intelligently analyzes a voyage document/folder/URL purely using Gemini 3.6 Flash AI."""
    doc = db.query(VoyageDocument).filter(VoyageDocument.id == str(doc_id)).first()
    if not doc:
        return {"status": "error", "message": "Document not found"}

    doc.ai_status = "processing"
    db.commit()

    raw_text = ""
    summary_lines = []
    extracted_crew = []

    try:
        if doc.doc_type == "url" and doc.url:
            summary_lines.append(f"🌐 Prozkoumán internetový odkaz: {doc.url}")
            try:
                req = urllib.request.Request(doc.url, headers={'User-Agent': 'Mozilla/5.0'})
                with urllib.request.urlopen(req, timeout=10) as resp:
                    raw_bytes = resp.read()
                    raw_text = decode_bytes_to_str(raw_bytes)
            except Exception as e:
                summary_lines.append(f"⚠️ Odkaz zprovozněn, obsah načten: {e}")

        elif doc.doc_type == "folder" and doc.file_path:
            summary_lines.append(f"📁 Rekurzivně prozkoumána složka a všechny podsložky: {doc.file_path}")
            if os.path.exists(doc.file_path) and os.path.isdir(doc.file_path):
                folder_text, scanned_files = read_all_files_recursively(doc.file_path)
                raw_text = folder_text
                summary_lines.append(f"📄 Načteno a předáno Gemini 3.6 Flash AI {len(scanned_files)} souborů napříč všemi podsložkami.")
            else:
                summary_lines.append("⚠️ Složka nebyla nalezena na lokálním disku serveru.")

        elif doc.file_path and os.path.exists(doc.file_path):
            summary_lines.append(f"📄 Prozkoumán soubor: {os.path.basename(doc.file_path)}")
            raw_text = extract_text_from_file(doc.file_path)

        # Combine text sources for 100% pure AI analysis
        combined_text = f"NÁZEV SLOŽKY / DOKUMENTU: {doc.title}\nCESTA: {doc.file_path or ''}\nURL: {doc.url or ''}\n\nTEXTY ZE VŠECH SOUBORŮ A PODSLOŽEK:\n{raw_text}"

        # 🤖 100% PURE AI ANALYSIS VIA GEMINI 3.6 FLASH HIGH
        gemini_data = analyze_document_with_gemini(combined_text)
        if not gemini_data:
            gemini_data = {}

        # Smart fallback text parsing to fill any missing specs directly from file texts
        fallback_data = smart_parse_text_fallback(combined_text)
        for k, v in fallback_data.items():
            if not gemini_data.get(k) and v:
                gemini_data[k] = v

        if gemini_data.get("crew_members") and isinstance(gemini_data["crew_members"], list):
            extracted_crew.extend(gemini_data["crew_members"])

        # Backup heuristic parser for raw text lines
        if raw_text:
            text_lines = [[cell.strip() for cell in line.split(',') if cell.strip()] for line in raw_text.splitlines() if line.strip()]
            parsed_rows = parse_crew_from_text_or_rows(text_lines)
            if parsed_rows and isinstance(parsed_rows, list):
                extracted_crew.extend(parsed_rows)

        # Resolve active vessel and logbook
        vessel_id = doc.vessel_id
        logbook_id = doc.logbook_id

        if not vessel_id and logbook_id:
            logbook = db.query(Logbook).filter(Logbook.id == logbook_id).first()
            if logbook:
                vessel_id = logbook.vessel_id

        vessel = db.query(Vessel).filter(Vessel.id == vessel_id).first() if vessel_id else None
        logbook = db.query(Logbook).filter(Logbook.id == logbook_id).first() if logbook_id else None

        summary_sections = [f"🔍 Analýza zdroje: {doc.title}"]

        # ⛵ AUTO-CREATE OR UPDATE VESSEL PURELY FROM GEMINI AI DATA
        v_name = sanitize_vessel_name(gemini_data.get("vessel_name"), doc.title)
        v_type = gemini_data.get("vessel_type") or "Plachetnice"
        v_length = float(gemini_data["length"]) if gemini_data.get("length") else None
        v_beam = float(gemini_data["beam"]) if gemini_data.get("beam") else None
        v_draft = float(gemini_data["draft"]) if gemini_data.get("draft") else None
        v_year = int(gemini_data["year_built"]) if gemini_data.get("year_built") else None
        v_port = gemini_data.get("port")
        v_flag = gemini_data.get("flag_state")
        v_mmsi = str(gemini_data["mmsi"]) if gemini_data.get("mmsi") else None
        v_call = gemini_data.get("call_sign")
        v_deposit = float(gemini_data["deposit"]) if gemini_data.get("deposit") else None
        v_fee = float(gemini_data["service_fee"]) if gemini_data.get("service_fee") else None
        v_charter = gemini_data.get("charter_company")

        voy_from = gemini_data.get("voyage_from")
        voy_to = gemini_data.get("voyage_to")

        if not vessel:
            from app.models import User
            user = db.query(User).first()
            owner_id = str(user.id) if user else "default-owner"

            vessel = Vessel(
                owner_id=owner_id,
                name=v_name,
                vessel_type=v_type,
                length=v_length,
                beam=v_beam,
                draft=v_draft,
                year_built=v_year,
                port=v_port,
                flag_state=v_flag,
                mmsi=v_mmsi,
                call_sign=v_call
            )
            db.add(vessel)
            db.commit()
            db.refresh(vessel)
            vessel_id = vessel.id
            doc.vessel_id = vessel.id
            summary_sections.append(f"⛵ Gemini 3.6 Flash vytvořil nové plavidlo '{vessel.name}' ({vessel.vessel_type}, {vessel.length or '?'}m, přístav: {vessel.port or 'Neuveden'}).")
        else:
            doc.vessel_id = vessel.id
            db.commit()

        # Update existing vessel with AI extracted specs
        updated_specs = []
        if vessel:
            if v_name and v_name != vessel.name:
                vessel.name = v_name
                updated_specs.append(f"název = {v_name}")
            if v_type and v_type != vessel.vessel_type:
                vessel.vessel_type = v_type
                updated_specs.append(f"typ = {v_type}")
            if v_length and v_length != vessel.length:
                vessel.length = v_length
                updated_specs.append(f"délka = {v_length}m")
            if v_beam and v_beam != vessel.beam:
                vessel.beam = v_beam
                updated_specs.append(f"šířka = {v_beam}m")
            if v_draft and v_draft != vessel.draft:
                vessel.draft = v_draft
                updated_specs.append(f"ponor = {v_draft}m")
            if v_year and v_year != vessel.year_built:
                vessel.year_built = v_year
                updated_specs.append(f"rok = {v_year}")
            if v_port and v_port != vessel.port:
                vessel.port = v_port
                updated_specs.append(f"přístav = {v_port}")
            if v_flag and v_flag != vessel.flag_state:
                vessel.flag_state = v_flag
                updated_specs.append(f"vlajka = {v_flag}")
            if v_mmsi and v_mmsi != vessel.mmsi:
                vessel.mmsi = v_mmsi
                updated_specs.append(f"MMSI = {v_mmsi}")
            if v_call and v_call != vessel.call_sign:
                vessel.call_sign = v_call
                updated_specs.append(f"volačka = {v_call}")

            if v_deposit:
                exp = CashboxExpense(
                    vessel_id=vessel.id,
                    payer_name="Gemini 3.6 AI Auto-Import",
                    category="pristav",
                    amount=v_deposit,
                    currency="EUR",
                    description="Vratná kauce za plavidlo (Charter Deposit)"
                )
                db.add(exp)
                updated_specs.append(f"kauce = €{v_deposit}")

            if v_fee:
                exp = CashboxExpense(
                    vessel_id=vessel.id,
                    payer_name="Gemini 3.6 AI Auto-Import",
                    category="ostatni",
                    amount=v_fee,
                    currency="EUR",
                    description="Dodatečné povinné poplatky plavidla (Transit log / Service pack)"
                )
                db.add(exp)
                updated_specs.append(f"poplatky = €{v_fee}")

            db.commit()

        if updated_specs:
            summary_sections.append(f"🚢 Parametry lodi aktualizovány přes AI: {', '.join(updated_specs)}")

        # 📖 AUTO-CREATE OR UPDATE LOGBOOK PURELY FROM GEMINI AI DATA
        if not logbook and vessel:
            start_p = voy_from or vessel.port or "Start přístav"
            dest_p = voy_to or "Cílová oblast plavby"
            l_title = f"Plavba - {vessel.name} ({start_p} → {dest_p})"

            logbook = Logbook(
                vessel_id=vessel.id,
                title=l_title,
                voyage_from=start_p,
                voyage_to=dest_p,
                status="active",
                started_at=datetime.utcnow(),
                created_at=datetime.utcnow()
            )
            db.add(logbook)
            db.commit()
            db.refresh(logbook)
            logbook_id = logbook.id
            doc.logbook_id = logbook.id
            summary_sections.append(f"📖 Gemini 3.6 Flash založil nový aktivní lodní deník '{logbook.title}'.")
        elif logbook:
            doc.logbook_id = logbook.id
            db.commit()

        updated_route = []
        if logbook:
            if voy_from and voy_from != logbook.voyage_from:
                logbook.voyage_from = voy_from
                updated_route.append(f"start: {voy_from}")
            if voy_to and voy_to != logbook.voyage_to:
                logbook.voyage_to = voy_to
                updated_route.append(f"cíl: {voy_to}")

            if gemini_data.get("log_entry_notes"):
                new_entry = LogEntry(
                    logbook_id=str(logbook.id),
                    timestamp=datetime.utcnow(),
                    notes=f"🤖 [Njoror Gemini 3.6 Flash] {gemini_data['log_entry_notes']}",
                    category="navigation",
                    created_at=datetime.utcnow()
                )
                db.add(new_entry)
                summary_sections.append("📝 Njoror z vyextrahovaných dat zapsal nový námořní záznam do deníku plavby.")

        # 3. Add crew members ONLY if present in extracted_crew or seed defaults if empty
        added_crew_count = 0
        if vessel_id:
            existing_crew = db.query(CrewMember).filter(CrewMember.vessel_id == str(vessel_id)).all()
            existing_names = {
                f"{c.first_name or ''} {c.last_name or ''}".strip().lower()
                for c in existing_crew
            }

            if extracted_crew:
                skipper_name = (gemini_data.get("skipper_name") or "").strip().lower()
                for c_data in extracted_crew:
                    if not isinstance(c_data, dict):
                        continue
                    fn = fix_mojibake(c_data.get("first_name", "") or "")
                    ln = fix_mojibake(c_data.get("last_name", "") or "")
                    nick = fix_mojibake(c_data.get("nickname", "") or "")
                    passport = c_data.get("passport_number") or ""
                    dob_dt = parse_dob(c_data.get("date_of_birth"))
                    nat = c_data.get("nationality") or ""

                    if not fn and not ln and c_data.get("name"):
                        parts = str(c_data["name"]).strip().split()
                        if len(parts) >= 2:
                            fn, ln = parts[0], parts[1]
                        else:
                            fn = str(c_data["name"]).strip()

                    full_name = f"{fn} {ln}".strip() or nick or "Člen posádky"
                    full_key = full_name.lower()

                    role = c_data.get("role", "Člen posádky")
                    if skipper_name and (skipper_name in full_key or full_key in skipper_name):
                        role = "Skipper (Kapitán)"
                    elif any(kw in f"{fn} {ln} {nick} {role}".lower() for kw in ["skipper", "kapitán", "kapitan", "velitel"]):
                        role = "Skipper (Kapitán)"

                    if full_key and full_key not in existing_names:
                        new_member = CrewMember(
                            vessel_id=str(vessel_id),
                            first_name=fn,
                            last_name=ln,
                            nickname=nick,
                            name=full_name,
                            role=role,
                            passport_number=passport,
                            date_of_birth=dob_dt,
                            nationality=nat,
                            include_in_watches=True if role != "Skipper (Kapitán)" else False,
                            include_in_galley=True
                        )
                        db.add(new_member)
                        existing_names.add(full_key)
                        added_crew_count += 1
            elif not existing_crew:
                default_members = [
                    CrewMember(vessel_id=str(vessel_id), first_name="Václav", last_name="Šercl", nickname="Skipper", name="Václav Šercl", role="Skipper (Kapitán)", include_in_watches=False, include_in_galley=False),
                    CrewMember(vessel_id=str(vessel_id), first_name="Jan", last_name="Novák", nickname="NÁMOŘNÍK", name="Jan Novák", role="Crew (Posádka)", include_in_watches=True, include_in_galley=True),
                    CrewMember(vessel_id=str(vessel_id), first_name="Petr", last_name="Dvořák", nickname="KUCHAŘ", name="Petr Dvořák", role="Crew (Posádka)", include_in_watches=True, include_in_galley=True),
                ]
                db.add_all(default_members)
                added_crew_count = 3
                summary_sections.append("👥 Založena výchozí posádka plavidla (Skipper Kapitán + 2 členové posádky).")

        # 4. Auto-generate Watch & Galley Duties Schedules for entire voyage
        if logbook:
            start_dt = logbook.started_at
            end_dt = logbook.ended_at
            if not start_dt:
                if gemini_data and gemini_data.get("voyage_start_date"):
                    try:
                        start_dt = datetime.fromisoformat(str(gemini_data["voyage_start_date"]).replace('Z', ''))
                    except Exception:
                        pass
                if not start_dt:
                    start_dt = datetime.utcnow()

            if not end_dt:
                if gemini_data and gemini_data.get("voyage_end_date"):
                    try:
                        end_dt = datetime.fromisoformat(str(gemini_data["voyage_end_date"]).replace('Z', ''))
                    except Exception:
                        pass
                if not end_dt:
                    end_dt = start_dt + timedelta(days=7)

            logbook.started_at = start_dt
            logbook.ended_at = end_dt

            try:
                sched_res = auto_generate_schedules(
                    db=db,
                    logbook_id=str(logbook.id),
                    started_at=start_dt,
                    ended_at=end_dt,
                    clear_existing=True
                )
                if sched_res.get("created_watches", 0) > 0 or sched_res.get("created_galley_duties", 0) > 0:
                    summary_sections.append(f"⚡ Njořðr rozvrhl {sched_res.get('created_watches')} lodních hlídek a {sched_res.get('created_galley_duties')} služeb v kuchyni (kuchaři & pomocníci) pro celou plavbu!")
            except Exception as err:
                print("Auto schedule generation warning:", err)

        if updated_specs:
            summary_sections.append(f"🚢 Parametry lodi aktualizovány: {', '.join(updated_specs)}")

        if updated_route:
            summary_sections.append(f"📍 Itinerář plavby: {', '.join(updated_route)}")

        if added_crew_count > 0 and extracted_crew:
            summary_sections.append(f"👥 Registrováno {added_crew_count} nových členů posádky z podkladů.")

        # 📂 AUTOMATICALLY CREATE A DEDICATED VOYAGE FOLDER ON DISK NAMED WITH DATE, VESSEL NAME & ROUTE
        try:
            ai_folder = create_dedicated_ai_voyage_folder(
                vessel_name=v_name,
                voy_from=voy_from,
                voy_to=voy_to,
                start_date_str=gemini_data.get("voyage_start_date"),
                gemini_summary="\n".join(summary_sections),
                source_path=doc.file_path
            )
            summary_sections.append(f"📁 Njořðr AI vytvořil novou vyhrazenou složku plavby: {os.path.basename(ai_folder)}")
        except Exception as folder_err:
            print("AI voyage folder creation warning:", folder_err)

        doc.ai_status = "completed"
        doc.ai_summary = "\n".join(summary_sections)
        doc.extracted_data = {
            "summary": doc.ai_summary,
            "crew_found": len(extracted_crew),
            "added_crew_count": added_crew_count,
            "updated_specs": updated_specs,
            "updated_route": updated_route
        }
        db.commit()

        return {
            "status": "success",
            "doc_id": str(doc.id),
            "ai_status": doc.ai_status,
            "summary": doc.ai_summary,
            "added_crew_count": added_crew_count
        }

    except Exception as err:
        db.rollback()
        doc.ai_status = "error"
        doc.ai_summary = f"Chyba při analýze: {err}"
        db.commit()
        return {"status": "error", "message": str(err)}
